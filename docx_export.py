"""Export daily recap to a DOCX file."""
import os
import re
from datetime import datetime
from zoneinfo import ZoneInfo
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()

TZ = ZoneInfo(os.getenv("TIMEZONE", "Asia/Dubai"))
RECAPS_DIR = os.path.join(os.path.dirname(__file__), "recaps")


def save_recap_docx(recap_text: str, date_str: str = None) -> str:
    """Save recap text as a DOCX file. Returns the file path."""
    os.makedirs(RECAPS_DIR, exist_ok=True)

    if not date_str:
        date_str = datetime.now(TZ).strftime("%d.%m.%Y")

    filename = f"recap_{date_str.replace('.', '-')}.docx"
    filepath = os.path.join(RECAPS_DIR, filename)

    doc = Document()

    # Title
    title = doc.add_heading(f"Daily Recap — {date_str}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Parse markdown-like recap text into formatted paragraphs
    for line in recap_text.split("\n"):
        line = line.rstrip()

        if not line:
            doc.add_paragraph("")
            continue

        # Headers
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("---"):
            # Horizontal rule — add a thin paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("—" * 40)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.size = Pt(8)
        elif line.startswith("- "):
            # Bullet point
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, line)

    doc.save(filepath)
    return filepath


def _add_formatted_text(paragraph, text: str):
    """Parse bold (**text**) and add runs to paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
